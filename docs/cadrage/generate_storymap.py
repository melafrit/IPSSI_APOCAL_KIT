#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_storymap.py — Genere directement la Story Map d'EduTutor IA en .docx.

Artefact de cadrage : User Story Mapping (Jeff Patton). La Story Map etant une
grille 2D, on la genere directement en Word (paysage), sans .md intermediaire.

Perimetre STRICT de l'exercice (chapitre 07) :
  1. Backbone (persona / epic / activite)
  2. Stories sous chaque activite (>= 3 par colonne)
  3. Decoupe MoSCoW en releases (Must / Should / Could)
  4. Walking Skeleton (MVP de bout en bout)
  5. Decompte par release + estimation Sprints 1-2

100 % trace aux artefacts de l'equipe : seuls figurent les codes US reels du
Product Backlog (E1-E12) et les stories enseignantes US-T.1/T.2/T.3 issues de la
perturbation J1 (documentees dans personas.md). Aucune story inventee, aucun
point d'estimation fabrique. L'etat d'avancement (livre/partiel/a faire) a ete
VERIFIE dans le code reel du depot (backend/ et frontend/), pas seulement repris
du backlog.

Dependance : python-docx     Usage : python generate_storymap.py
"""
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "story-map.docx")

INDIGO = RGBColor(0x4F, 0x46, 0xE5)
SLATE = RGBColor(0x33, 0x41, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SKEL = RGBColor(0xB4, 0x53, 0x09)

FILL_HEADER = "4F46E5"
FILL_MUST_LABEL = "059669"
FILL_MUST = "D1FAE5"
FILL_SHOULD_LABEL = "D97706"
FILL_SHOULD = "FEF3C7"
FILL_COULD_LABEL = "475569"
FILL_COULD = "F1F5F9"

# ---------------------------------------------------------------------------
# DONNEES — uniquement des US reelles du backlog + US-T.x (perturbation J1).
# Story = (code, label_court, label_complet, points, etat)
# points "-" = non estime au backlog ; etat : livre / partiel / vide
# Etats verifies dans le code : backend/{accounts,quizzes,llm,administration},
# frontend/src/pages, .github/workflows, settings.py.
# ---------------------------------------------------------------------------
COLUMNS = [
    {
        "persona": "Lucas Moreau",
        "epic": "Compte & identite (E1)",
        "activite": "Creer & gerer son compte",
        "must": [
            ("US-F1.1", "Creer un compte", "Creer un compte par email", "2", "livre"),
            ("US-F1.2", "Login / Logout", "Se connecter / se deconnecter", "2", "livre"),
            ("US-F1.3", "Valider l'email", "Valider son email via un lien (+ renvoi du lien)", "3", "livre"),
            ("US-F1.4", "Reset mot de passe", "Reinitialiser son mot de passe oublie", "3", "livre"),
            ("US-F1.5", "Profil", "Consulter, modifier et supprimer son profil", "2", "livre"),
        ],
        "should": [
            ("US-F1.6", "Erreurs claires", "Messages d'erreur clairs (email pris, identifiants invalides)", "2", "partiel"),
        ],
        "could": [],
    },
    {
        "persona": "Lucas Moreau",
        "epic": "Ingestion de cours (E2)",
        "activite": "Deposer un cours",
        "must": [
            ("US-F2.1", "Upload PDF <=5Mo", "Televerser un PDF (<= 5 Mo) comme source de quiz", "5", "livre"),
            ("US-F2.2", "Coller du texte", "Coller du texte (>= 200 car.) comme source de quiz", "2", "livre"),
            ("US-F2.3", "Erreur fichier/texte", "Message d'erreur si fichier trop lourd / texte trop court", "2", "livre"),
        ],
        "should": [
            ("US-F2.4", "Alerte extraction", "Prevenir si l'extraction PDF echoue (PDF scanne/image)", "2", "livre"),
        ],
        "could": [],
    },
    {
        "persona": "Lucas Moreau",
        "epic": "Generation & passation (E3/E4)",
        "activite": "Generer & passer le quiz",
        "must": [
            ("US-F3.1", "Generer 10 QCM", "Generer 10 QCM via le LLM local, ancres dans le cours", "5", "livre"),
            ("US-F3.2", "4 options / 1 bonne", "Chaque QCM a 4 options et 1 bonne reponse", "3", "livre"),
            ("US-F3.4", "Erreur LLM / timeout", "Message exploitable si le LLM echoue / time-out", "3", "livre"),
            ("US-F4.1", "Soumettre", "Soumettre ses reponses pour correction automatique", "3", "livre"),
            ("US-F4.2", "Score juste", "Une seule bonne reponse comptee par QCM", "2", "livre"),
        ],
        "should": [
            ("US-F3.3", "Chargement", "Retour de chargement pendant la generation", "2", "livre"),
            ("US-F3.5", "Ancrage RAG", "QCM reellement ancres dans le cours (RAG)", "5", "partiel"),
        ],
        "could": [
            ("US-R2.1", "RAG renforce", "Generation ancree chapitre par chapitre", "8", "partiel"),
        ],
    },
    {
        "persona": "Lucas Moreau",
        "epic": "Resultats (E5)",
        "activite": "Voir ses resultats",
        "must": [
            ("US-F5.1", "Score /10", "Voir son score /10 apres soumission", "2", "livre"),
            ("US-F5.2", "Detail des reponses", "Detail bonne/mauvaise reponse par question", "3", "livre"),
        ],
        "should": [
            ("US-F5.3", "Page d'accueil", "Page d'accueil claire expliquant le produit", "2", "livre"),
        ],
        "could": [],
    },
    {
        "persona": "Lucas Moreau",
        "epic": "Historique (E6)",
        "activite": "Suivre sa progression",
        "must": [
            ("US-F6.1", "Historique", "Retrouver l'historique de ses quiz (date, cours, score)", "3", "livre"),
        ],
        "should": [
            ("US-F6.2", "Rouvrir un quiz", "Rouvrir un quiz passe pour revoir ses reponses", "2", "livre"),
        ],
        "could": [
            ("US-F6.3", "Filtrer / trier", "Filtrer/trier l'historique (date, score, matiere)", "3", "vide"),
        ],
    },
    {
        "persona": "Mme Lefevre (J1)",
        "epic": "Espace enseignant (J1)",
        "activite": "Piloter sa classe",
        "must": [
            ("US-T.1", "Dashboard classe", "Voir scores & progression de ses 28 etudiants", "-", "vide"),
            ("US-T.2", "Reperer decrocheurs", "Tri/alerte sur scores faibles en <= 3 clics", "-", "vide"),
            ("US-T.3", "Conseils cibles", "Envoyer un conseil cible a un etudiant en difficulte", "-", "vide"),
        ],
        "should": [
            ("US-R2.2", "Export / partage", "Tableau de bord enseignant : suivi de classe, export", "8", "vide"),
        ],
        "could": [],
    },
    {
        "persona": "Equipe / Admin",
        "epic": "Securite & RGPD (E7/E8)",
        "activite": "Securiser & se conformer",
        "must": [
            ("US-S.1", "Anti prompt-injection", "Separer instructions systeme et contenu utilisateur", "5", "partiel"),
            ("US-S.2", "Assainir les entrees", "Valider/assainir les entrees (cours, texte)", "3", "partiel"),
            ("US-S.3", "Isoler les donnees", "Isoler les donnees par utilisateur", "3", "livre"),
            ("US-S.6", "Pas de secrets", "Ne jamais exposer les secrets", "2", "partiel"),
            ("US-G.1", "Confidentialite", "Politique de confidentialite", "3", "vide"),
            ("US-G.2", "Droit a l'oubli", "Supprimer definitivement compte et donnees", "3", "livre"),
        ],
        "should": [
            ("US-S.4", "Tests adversariaux", "Tests adversariaux d'injection", "3", "vide"),
            ("US-S.5", "Rate limiting", "Rate limiting / quotas sur la generation", "3", "vide"),
            ("US-G.3", "CGU / mentions", "CGU et mentions legales", "2", "vide"),
            ("US-G.4", "Consentement", "Gestion du consentement (cookies/tracage)", "3", "vide"),
            ("US-G.6", "Minimisation", "Minimiser et documenter les donnees collectees", "2", "vide"),
        ],
        "could": [
            ("US-G.5", "Portabilite", "Exporter ses donnees (portabilite RGPD)", "3", "vide"),
        ],
    },
    {
        "persona": "Equipe / Admin",
        "epic": "Livraison & qualite (E9/E10)",
        "activite": "Livrer & exploiter",
        "must": [
            ("US-D.1", "CI verte", "CI verte (lint + tests) a chaque PR", "3", "livre"),
            ("US-D.2", "Lancement 1 commande", "Lancement en une commande (scripts par OS)", "2", "livre"),
            ("US-D.3", "Plan de crise", "Plan de crise / rollback pour le jour J", "3", "vide"),
            ("US-X.3", "ADR choix LLM", "Tracer le choix du fournisseur LLM (ADR)", "3", "vide"),
        ],
        "should": [
            ("US-X.1", "Tests F3/F4", "Tests automatises sur le coeur metier (F3/F4)", "5", "livre"),
            ("US-X.2", "Mock / fallback LLM", "Mode mock/fallback LLM pour securiser la demo", "3", "livre"),
            ("US-X.4", "Config LLM via UI", "Configurer le LLM/l'app depuis l'UI (admin)", "5", "livre"),
            ("US-D.5", "Post-mortem", "Post-mortem apres la crise", "1", "vide"),
        ],
        "could": [
            ("US-D.4", "Deploiement prod", "Deploiement de production (VPS OVH, HTTPS)", "8", "partiel"),
            ("US-X.5", "Observabilite", "Logs / observabilite (erreurs LLM, latence)", "3", "vide"),
        ],
    },
]

SKELETON = {"US-F1.1", "US-F2.1", "US-F3.1", "US-F4.1", "US-F5.1", "US-F6.1"}
ETAT_LABEL = {"livre": "OK livre", "partiel": "Partiel", "vide": "A faire"}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_text(cell, text, size=8, bold=False, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color


def fill_story_cell(cell, stories, fill):
    shade(cell, fill)
    cell.text = ""
    if not stories:
        set_cell_text(cell, "—", size=8, color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    first = True
    for code, short, _full, _pts, _etat in stories:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        star = "★ " if code in SKELETON else ""
        rc = p.add_run(f"{star}{code}")
        rc.bold = True
        rc.font.size = Pt(7)
        rc.font.color.rgb = SKEL if star else INDIGO
        rl = p.add_run(f"  {short}")
        rl.font.size = Pt(7)
        rl.font.color.rgb = SLATE


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------
def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = sec.right_margin = Cm(1.4)
    sec.top_margin = sec.bottom_margin = Cm(1.4)

    cover_page(doc)
    intro(doc)
    backbone_table(doc)
    story_map_grid(doc)
    walking_skeleton(doc)
    detail_tables(doc)
    analysis(doc)

    doc.save(OUT)
    return OUT


def cover_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("EduTutor IA"); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = INDIGO
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cadrage produit"); r.font.size = Pt(16); r.font.color.rgb = SLATE
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Story Map — User Story Mapping (Jeff Patton)"); r.bold = True; r.font.size = Pt(22)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Livrable Jour 1 — APOCAL'IPSSI 2026 — Equipe 6"); r.italic = True; r.font.color.rgb = SLATE
    doc.add_page_break()


def intro(doc):
    doc.add_heading("1. Objectif & methode", level=1)
    doc.add_paragraph(
        "Cette Story Map organise le perimetre d'EduTutor IA selon la technique du User Story "
        "Mapping (Jeff Patton). Elle transforme le Product Backlog « plat » en carte a deux dimensions :"
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Axe horizontal (->) : ").bold = True
    p.add_run("le parcours utilisateur dans le temps (les activites, de gauche a droite).")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Axe vertical (v) : ").bold = True
    p.add_run("la priorite MoSCoW, des stories essentielles (Release 1) aux optionnelles (Release 3).")
    doc.add_paragraph(
        "Elle ne reorganise que l'existant : toutes les stories sont des codes US du Product Backlog "
        "de l'equipe (epics E1-E12), plus les stories enseignantes US-T.1/T.2/T.3 issues de la "
        "perturbation J1. Le parcours horizontal suit Lucas Moreau, persona etudiant primaire du "
        "Customer Journey Map ; Mme Lefevre porte la branche enseignante issue de la perturbation J1."
    )
    p = doc.add_paragraph()
    p.add_run("Etat d'avancement reel : ").bold = True
    p.add_run(
        "la colonne « Etat code » a ete verifiee directement dans le code du depot (backend Django : "
        "accounts, quizzes, llm, administration ; frontend React ; CI). Elle ne se contente pas de recopier "
        "le backlog. Valeurs : OK livre (implemente et cable), Partiel (present mais a fiabiliser/finir), "
        "A faire (absent)."
    )
    p = doc.add_paragraph()
    p.add_run("Note perturbation J1 : ").bold = True
    p.add_run(
        "le PO a repositionne Mme Lefevre en persona primaire (produit enseignant-first). La branche "
        "« Piloter sa classe » (US-T.1/T.2/T.3) passe donc en Release 1 ; ces stories restent a estimer "
        "au backlog et sont aujourd'hui entierement a faire (aucun code d'espace enseignant present)."
    )


def backbone_table(doc):
    doc.add_heading("2. Le Backbone (colonne vertebrale)", level=1)
    doc.add_paragraph(
        "Le backbone resume le parcours en une ligne : qui (persona), sur quelle thematique (epic), "
        "fait quelle action (activite)."
    )
    table = doc.add_table(rows=3, cols=len(COLUMNS) + 1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(table.rows[0].cells[0], "Persona", size=8, bold=True, color=WHITE)
    set_cell_text(table.rows[1].cells[0], "Epic", size=8, bold=True, color=WHITE)
    set_cell_text(table.rows[2].cells[0], "Activite", size=8, bold=True, color=WHITE)
    for i in range(3):
        shade(table.rows[i].cells[0], FILL_HEADER)
    for j, col in enumerate(COLUMNS, start=1):
        set_cell_text(table.rows[0].cells[j], col["persona"], size=8, bold=True, color=INDIGO, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.rows[1].cells[j], col["epic"], size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.rows[2].cells[j], col["activite"], size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(table.rows[2].cells[j], FILL_COULD)
    doc.add_paragraph()


def story_map_grid(doc):
    doc.add_heading("3. La Story Map (parcours x releases)", level=1)
    doc.add_paragraph(
        "Lecture horizontale : le parcours complet. Lecture verticale : les stories par priorite MoSCoW, "
        "decoupees en 3 releases. Les stories marquees ★ composent le Walking Skeleton (section 4)."
    )
    ncols = len(COLUMNS) + 1
    table = doc.add_table(rows=4, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_cell_text(table.rows[0].cells[0], "", size=8)
    shade(table.rows[0].cells[0], FILL_HEADER)
    for j, col in enumerate(COLUMNS, start=1):
        set_cell_text(table.rows[0].cells[j], col["activite"], size=7, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(table.rows[0].cells[j], FILL_HEADER)

    bands = [
        (1, "RELEASE 1\nMUST\n(MVP)", FILL_MUST_LABEL, FILL_MUST, "must"),
        (2, "RELEASE 2\nSHOULD", FILL_SHOULD_LABEL, FILL_SHOULD, "should"),
        (3, "RELEASE 3\nCOULD", FILL_COULD_LABEL, FILL_COULD, "could"),
    ]
    for row_idx, label, label_fill, cell_fill, key in bands:
        lab = table.rows[row_idx].cells[0]
        set_cell_text(lab, label, size=8, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(lab, label_fill)
        for j, col in enumerate(COLUMNS, start=1):
            fill_story_cell(table.rows[row_idx].cells[j], col[key], cell_fill)

    for r in table.rows:
        r.cells[0].width = Cm(1.9)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Legende : ").bold = True
    p.add_run("★ = story du Walking Skeleton (MVP de bout en bout).")


def walking_skeleton(doc):
    doc.add_heading("4. Le Walking Skeleton (MVP minimal de bout en bout)", level=1)
    doc.add_paragraph(
        "Le Walking Skeleton est la plus petite suite de stories qui permet a Lucas de traverser tout le "
        "parcours. C'est l'objectif de la Release 1 du backlog : "
        "« je m'inscris -> je depose un cours -> je revise -> je vois mon score -> je retrouve mon historique »."
    )
    steps = [
        ("US-F1.1", "Je cree mon compte"),
        ("US-F2.1", "Je depose un cours (PDF)"),
        ("US-F3.1", "Je genere 10 QCM"),
        ("US-F4.1", "Je reponds et je soumets"),
        ("US-F5.1", "Je vois mon score /10"),
        ("US-F6.1", "Je retrouve mon historique"),
    ]
    p = doc.add_paragraph()
    for i, (code, label) in enumerate(steps):
        r = p.add_run(code); r.bold = True; r.font.color.rgb = SKEL
        r2 = p.add_run(f" {label}"); r2.font.color.rgb = SLATE
        if i < len(steps) - 1:
            sep = p.add_run("   ->   "); sep.bold = True; sep.font.color.rgb = INDIGO
    doc.add_paragraph(
        "Bonne nouvelle : ce squelette est deja entierement livre dans le code (le flux fonctionne de bout "
        "en bout). L'effort de la semaine porte donc sur le durcissement, la conformite et la branche "
        "enseignante, pas sur le flux nominal."
    )


def detail_tables(doc):
    doc.add_heading("5. Detail des stories par activite", level=1)
    doc.add_paragraph(
        "Pour chaque activite : les user stories avec priorite MoSCoW, estimation en points (Fibonacci, "
        "telle qu'au backlog) et etat verifie dans le code."
    )
    rel = {"must": "R1 - Must", "should": "R2 - Should", "could": "R3 - Could"}
    for col in COLUMNS:
        doc.add_heading(f"{col['activite']}  ({col['epic']})", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Code", "User Story", "Priorite", "Pts", "Etat code"]):
            set_cell_text(table.rows[0].cells[i], h, size=9, bold=True)
        for key in ("must", "should", "could"):
            for code, _short, full, pts, etat in col[key]:
                c = table.add_row().cells
                set_cell_text(c[0], code, size=9, bold=True)
                set_cell_text(c[1], full, size=9)
                set_cell_text(c[2], rel[key], size=9)
                set_cell_text(c[3], pts, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(c[4], ETAT_LABEL[etat], size=9)
        doc.add_paragraph()


def _release_stats(key):
    pts = count = unestimated = 0
    etats = {"livre": 0, "partiel": 0, "vide": 0}
    for col in COLUMNS:
        for _code, _short, _full, p, etat in col[key]:
            count += 1
            etats[etat] += 1
            if p.isdigit():
                pts += int(p)
            else:
                unestimated += 1
    return pts, count, unestimated, etats


def analysis(doc):
    doc.add_heading("6. Decompte & tenue dans les Sprints 1-2", level=1)

    r1_pts, r1_cnt, r1_un, r1_e = _release_stats("must")
    r2_pts, r2_cnt, _, _ = _release_stats("should")
    r3_pts, r3_cnt, _, _ = _release_stats("could")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Release", "Priorite", "Nb stories", "Points estimes", "Avancement (livre/partiel/a faire)"]):
        set_cell_text(table.rows[0].cells[i], h, size=10, bold=True)
    _, _, _, e2 = _release_stats("should")
    _, _, _, e3 = _release_stats("could")
    rows = [
        ("Release 1 (MVP)", "Must", r1_cnt, f"{r1_pts} (+{r1_un} a estimer)", r1_e),
        ("Release 2", "Should", r2_cnt, str(r2_pts), e2),
        ("Release 3", "Could", r3_cnt, str(r3_pts), e3),
    ]
    for rel_name, prio, cnt, pts, e in rows:
        c = table.add_row().cells
        set_cell_text(c[0], rel_name, size=10)
        set_cell_text(c[1], prio, size=10)
        set_cell_text(c[2], str(cnt), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(c[3], pts, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(c[4], f"{e['livre']} / {e['partiel']} / {e['vide']}", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    doc.add_heading("La Release 1 tient-elle dans les Sprints 1-2 ?", level=2)
    doc.add_paragraph(
        f"La Release 1 compte {r1_cnt} stories Must pour ~{r1_pts} points estimes ({r1_un} stories "
        f"enseignantes restant a estimer). Surtout, l'audit du code montre que {r1_e['livre']} sont deja "
        f"livrees, {r1_e['partiel']} partielles et {r1_e['vide']} a faire."
    )
    doc.add_paragraph(
        "Le flux nominal (Walking Skeleton + coeur etudiant E1 a E6) est deja fonctionnel : l'effort "
        "Sprints 1-2 se concentre sur le fiabiliser et le tester, pas sur le construire.",
        style="List Bullet",
    )
    doc.add_paragraph(
        f"Les {r1_e['vide']} stories a faire sont la branche enseignante (J1) et le socle "
        "conformite/securite/crise, cadences par les perturbations J3 a J5 et etales sur la semaine.",
        style="List Bullet",
    )
    p = doc.add_paragraph()
    p.add_run("Estimation : ").bold = True
    p.add_run(
        "oui, la Release 1 tient a l'echelle de la semaine (R1 due mercredi 17h45). Le coeur etudiant etant "
        "deja livre, les Sprints 1-2 absorbent le durcissement du MVP ; la branche enseignante et la "
        "conformite arrivent ensuite, au rythme des perturbations."
    )


if __name__ == "__main__":
    out = build()
    print(f"[OK] {os.path.relpath(out, HERE)} genere.")
